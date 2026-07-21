from pathlib import Path

from docx import Document


OUTPUT_FOLDER = Path("generated_reports")

OUTPUT_FOLDER.mkdir(exist_ok=True)


def create_resume_docx(filename: str, content: str):

    path = OUTPUT_FOLDER / filename

    document = Document()

    document.add_heading("Rewritten Resume", level=1)

    document.add_paragraph(content)

    document.save(path)

    return path